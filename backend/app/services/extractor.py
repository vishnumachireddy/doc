import os
import docx
from pypdf import PdfReader
import pdfplumber
from concurrent.futures import ThreadPoolExecutor
from typing import List, Tuple

def extract_pdf_page(pdf_path: str, page_num: int) -> Tuple[int, str]:
    """Extract text from a single PDF page for multi-threaded extraction."""
    try:
        with pdfplumber.open(pdf_path) as pdf:
            if page_num < len(pdf.pages):
                page_text = pdf.pages[page_num].extract_text()
                return page_num, page_text or ""
    except Exception as e:
        print(f"Error extracting page {page_num} of {pdf_path}: {e}")
    return page_num, ""

def extract_pdf_text_parallel(pdf_path: str) -> Tuple[str, int]:
    """
    Extracts PDF text by dividing pages and running text extraction in parallel
    using a thread pool. Returns the full text and page count.
    """
    try:
        reader = PdfReader(pdf_path)
        num_pages = len(reader.pages)
        if num_pages == 0:
            return "", 0
        
        # Parallel extraction of pages
        results = [None] * num_pages
        with ThreadPoolExecutor(max_workers=min(4, num_pages)) as executor:
            future_to_page = {
                executor.submit(extract_pdf_page, pdf_path, page_num): page_num 
                for page_num in range(num_pages)
            }
            for future in future_to_page:
                page_num, text = future.result()
                results[page_num] = text

        full_text = "\n--- PAGE BREAK ---\n".join(results)
        return full_text, num_pages
    except Exception as e:
        print(f"Error in parallel PDF extraction: {e}")
        # Fallback to simple extraction
        try:
            reader = PdfReader(pdf_path)
            pages_text = [page.extract_text() or "" for page in reader.pages]
            return "\n--- PAGE BREAK ---\n".join(pages_text), len(reader.pages)
        except Exception:
            return "", 0

def extract_docx_text(docx_path: str) -> str:
    """Extracts text from a DOCX document including paragraphs and tables."""
    try:
        doc = docx.Document(docx_path)
        full_text = []
        
        # Extract from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
                
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    full_text.append(" | ".join(row_text))
                    
        return "\n".join(full_text)
    except Exception as e:
        print(f"Error extracting DOCX text: {e}")
        return ""

def extract_document_text(file_path: str) -> Tuple[str, int]:
    """
    Main extraction router. Supported formats: pdf, docx.
    Returns (extracted_text, page_count).
    Images are handled by OCR separately, so they return empty string and 1 page here.
    """
    ext = os.path.splitext(file_path.lower())[1].lstrip(".")
    if ext == "pdf":
        return extract_pdf_text_parallel(file_path)
    elif ext == "docx":
        text = extract_docx_text(file_path)
        return text, 1
    elif ext in ["png", "jpg", "jpeg"]:
        # Images are 1 page; text will be extracted by OCR later
        return "", 1
    return "", 0
