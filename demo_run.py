import os
import docx
import uuid
from backend.app.core.database import SessionLocal, Base, engine
from backend.app.core.security import get_password_hash, verify_password
from backend.app.models.user import User
from backend.app.models.document import Document
from backend.app.models.analysis import ResumeAnalysis
from backend.app.worker.tasks import process_document_core
from backend.app.services.job_matcher import match_resume_to_jd

def run_demo():
    print("="*60)
    print("             RESUMEIQ END-TO-END DEMONSTRATION")
    print("="*60)
    
    # 1. Create a real sample resume Word Document
    os.makedirs("uploads", exist_ok=True)
    doc_path = "uploads/demo_resume.docx"
    
    doc = docx.Document()
    doc.add_heading("John Doe", level=0)
    doc.add_paragraph("Email: john.doe@email.com | Phone: +91 9876543210")
    
    doc.add_heading("PROFESSIONAL EXPERIENCE", level=1)
    
    # Job 1 (2.4 years of Python/FastAPI/AWS)
    p1 = doc.add_paragraph()
    p1.add_run("Senior Software Engineer at TechCorp (Jan 2021 to Present)\n").bold = True
    p1.add_run(
        "Spearheaded development of a high-throughput API gateway using Python, FastAPI, and PostgreSQL. "
        "Leveraged AWS and Docker to orchestrate containerized microservices. Fostered collaborative agile teamwork "
        "and mentored junior engineers. Optimized database query performance, reducing latency by 45%."
    )
    
    # Job 2 (2.5 years of React/JavaScript/MySQL)
    p2 = doc.add_paragraph()
    p2.add_run("Software Engineer at AppDev Solutions (June 2018 to Dec 2020)\n").bold = True
    p2.add_run(
        "Developed web applications using JavaScript, React, and Python. Configured databases using SQL and MySQL. "
        "Maintained version control systems with Git. Participated in project management cycles."
    )
    
    doc.add_heading("EDUCATION", level=1)
    doc.add_paragraph("Bachelor of Technology in Computer Science, Engineering Institute (2018)")
    
    doc.add_heading("TECHNICAL SKILLS", level=1)
    doc.add_paragraph("Python, JavaScript, FastAPI, React, PostgreSQL, MySQL, Docker, AWS, Git, Project Management")
    
    doc.save(doc_path)
    print(f"[+] Step 1: Created sample resume file at: {doc_path}")

    # 2. Database setup
    db = SessionLocal()
    try:
        # Delete previous demo records to ensure clean run
        db.query(User).filter(User.email == "demo@gmail.com").delete()
        db.commit()

        # Create User account
        hashed = get_password_hash("password123")
        user = User(email="demo@gmail.com", hashed_password=hashed)
        db.add(user)
        db.commit()
        db.refresh(user)
        print(f"[+] Step 2: Registered user: {user.email} (Password: password123)")

        # 3. Create Document Record
        doc_id = uuid.uuid4().hex
        db_doc = Document(
            id=doc_id,
            user_id=user.id,
            file_path=doc_path,
            file_hash="demo_hash_placeholder",
            file_size=os.path.getsize(doc_path),
            mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            status="pending"
        )
        db.add(db_doc)
        db.commit()
        db.refresh(db_doc)
        print(f"[+] Step 3: Created document metadata (ID: {db_doc.id})")

        # 4. Run Processing Pipeline Synchronously
        print("\n[+] Step 4: Running AI Processing Pipeline (Visual & Textual classification, Quality Verification, Redaction)...")
        success = process_document_core(db, doc_id)
        
        db.refresh(db_doc)
        print("\n" + "-"*40)
        print("          DOCUMENT METADATA RESULTS")
        print("-"*40)
        print(f"Processing Status:          {db_doc.status}")
        print(f"Classification Label:       {db_doc.classification_label}")
        print(f"Classification Confidence:  {int(db_doc.classification_confidence * 100)}%")
        print(f"Verification Quality Score: {db_doc.verification_score}/100")
        print(f"Verification Details:       {db_doc.verification_details}")
        print(f"Cryptographic Hash:         {db_doc.file_hash}")
        
        # 5. Fetch Resume Analysis
        analysis = db.query(ResumeAnalysis).filter(ResumeAnalysis.document_id == doc_id).first()
        if analysis:
            print("\n" + "-"*40)
            print("          RESUME INTELLIGENCE REPORT")
            print("-"*40)
            print(f"ATS Compliance Audit Score: {analysis.ats_score}/100")
            print(f"AI Assistance Estimate:     {analysis.ai_likelihood_score}%")
            print(f"AI Classification Label:    {analysis.suggestions.get('ai_classification')}")
            print(f"Recruiter Readiness Score:  {analysis.extracted_data.get('recruiter_readiness_score')}%")
            print(f"Extracted Contact:          {analysis.extracted_data.get('contact')}")
            print(f"Extracted Skills:           {analysis.extracted_data.get('skills')}")
            print(f"Skill Chronology (Years of Experience per skill):")
            for skill, years in analysis.extracted_data.get("skill_durations", {}).items():
                print(f"  - {skill}: {years} Years")
            
            # 6. Test Job Description Matching
            print("\n" + "-"*40)
            print("          SEMANTIC JOB MATCH REPORT")
            print("-"*40)
            target_jd = (
                "Looking for a Python Developer. Must have experience with FastAPI, PostgreSQL, and AWS. "
                "Familiarity with React is desired. Agile and project management skills are a plus."
            )
            print(f"Target Job Description:\n{target_jd}\n")
            
            match_res = match_resume_to_jd(db_doc.extracted_text, target_jd, analysis.extracted_data.get("skills", []))
            print(f"JD Match score:      {match_res['skill_match_percentage']}%")
            print(f"Missing Skills:      {match_res['missing_skills']}")
            print(f"Missing Keywords:    {match_res['missing_keywords']}")
            print(f"Suggestions:         {match_res['improvement_suggestions']}")
            
        print("\n" + "="*60)
        print("          DEMONSTRATION RUN COMPLETED SUCCESSFULLY")
        print("="*60)
        
    finally:
        db.close()

if __name__ == "__main__":
    run_demo()
